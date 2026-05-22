from docx import Document
from datetime import datetime
import os
import logging

# Asegurar logs
os.makedirs("logs", exist_ok=True)
logging.basicConfig(
    filename="logs/procesos.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

def generar_documento_word(datos: dict):
    """
    Genera documento Word basado en plantilla persona_natural.docx
    """

    if not datos:
        raise ValueError("El diccionario 'datos' está vacío")

    # 📁 Crear carpeta si no existe
    os.makedirs("generados", exist_ok=True)

    ruta_plantilla = os.path.join("plantillas", "persona_natural.docx")

    if not os.path.exists(ruta_plantilla):
        raise FileNotFoundError("No existe la plantilla persona_natural.docx")

    logging.info("Iniciando generación de documento Word")

    doc = Document(ruta_plantilla)

    # 🔥 MAPEO SEGURO (no rompe si falta algo)
    variables = {
        "{{nombre_propietario}}": str(datos.get("nombre_propietario", "")),
        "{{cedula_propietario}}": str(datos.get("cedula_propietario", "")),
        "{{direccion_predio}}": str(datos.get("direccion_predio", "")),
        "{{nombre_apoderado}}": str(datos.get("nombre_apoderado", "")),
        "{{cedula_apoderado}}": str(datos.get("cedula_apoderado", "")),
        "{{lugar_expedicion_apoderado}}": str(datos.get("lugar_expedicion_apoderado", "")),
    }

    # 🔥 REEMPLAZO ROBUSTO (no pierde formato)
    def reemplazar_texto(parrafo):
        for clave, valor in variables.items():
            if clave in parrafo.text:
                inline = parrafo.runs
                for i in range(len(inline)):
                    if clave in inline[i].text:
                        inline[i].text = inline[i].text.replace(clave, valor)

    # 📄 Párrafos normales
    for p in doc.paragraphs:
        reemplazar_texto(p)

    # 📊 Tablas (CRÍTICO)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    reemplazar_texto(p)

    # 🕒 Nombre único
    fecha = datetime.now().strftime("%Y%m%d_%H%M%S")
    nombre_archivo = f"poder_{datos.get('cedula_propietario','sincc')}_{fecha}.docx"
    ruta_salida = os.path.join("generados", nombre_archivo)

    # 💾 Guardar
    doc.save(ruta_salida)

    logging.info(f"Word generado correctamente: {ruta_salida}")

    return ruta_salida, nombre_archivo
